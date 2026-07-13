"""Test fixtures and shared configuration."""

import io
import os
import zipfile

import pytest
from fastapi.testclient import TestClient


@pytest.fixture
def test_client():
    """FastAPI test client with mocked model."""
    # Patch model loading before importing app
    from unittest.mock import patch

    with patch("app.services.ner_extractor.load_model"):
        from app.main import app

        yield TestClient(app)


@pytest.fixture
def sample_pdf_bytes():
    """Create a minimal valid PDF file for testing."""
    # Minimal valid PDF structure
    pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj

2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj

3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj

4 0 obj
<< /Length 178 >>
stream
BT
/F1 12 Tf
50 700 Td
(John Doe) Tj
0 -20 Td
(john.doe@email.com) Tj
0 -20 Td
(+1 555-123-4567) Tj
0 -20 Td
(Software Engineer at Google) Tj
0 -20 Td
(Skills: Python, Java, Docker) Tj
ET
endstream
endobj

5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj

xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000062 00000 n
0000000115 00000 n
0000000266 00000 n
0000000496 00000 n

trailer
<< /Root 1 0 R /Size 6 >>
startxref
573
%%EOF"""
    return pdf_content


@pytest.fixture
def sample_docx_bytes():
    """Create a minimal valid DOCX file for testing."""
    buf = io.BytesIO()
    from docx import Document

    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("john.doe@email.com")
    doc.add_paragraph("+1 555-123-4567")
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("Google - 2020 to 2024")
    doc.add_paragraph("Skills: Python, Java, Docker, Kubernetes")
    doc.add_paragraph("BS Computer Science - MIT - 2020")
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def empty_file_bytes():
    """Empty file (0 bytes)."""
    return b""


@pytest.fixture
def text_file_bytes():
    """A plain text file (invalid MIME type)."""
    return b"This is just a plain text file, not a CV."


@pytest.fixture
def large_file_bytes():
    """A file exceeding the max size limit (>10MB)."""
    # Create 11MB of data
    return b"A" * (11 * 1024 * 1024)


@pytest.fixture
def corrupt_pdf_bytes():
    """A corrupt/invalid PDF file."""
    return b"%PDF-1.4\nThis is not a valid PDF content.\n%%EOF"


@pytest.fixture
def mock_ner_entities():
    """Sample NER entity output for testing normalizer."""
    from app.schemas import NEREntity

    return [
        NEREntity(
            text="John Doe", label="name", score=0.95, start=0, end=8
        ),
        NEREntity(
            text="john.doe@email.com",
            label="email",
            score=0.98,
            start=9,
            end=27,
        ),
        NEREntity(
            text="+1 555-123-4567",
            label="phone",
            score=0.92,
            start=28,
            end=43,
        ),
        NEREntity(
            text="Software Engineer",
            label="title",
            score=0.88,
            start=44,
            end=61,
        ),
        NEREntity(
            text="Google",
            label="company",
            score=0.90,
            start=62,
            end=68,
        ),
        NEREntity(
            text="Python",
            label="skills",
            score=0.85,
            start=69,
            end=75,
        ),
        NEREntity(
            text="Java",
            label="skills",
            score=0.82,
            start=77,
            end=81,
        ),
        NEREntity(
            text="MIT",
            label="institution",
            score=0.91,
            start=82,
            end=85,
        ),
        NEREntity(
            text="BS Computer Science",
            label="degree",
            score=0.87,
            start=86,
            end=105,
        ),
        NEREntity(
            text="2020",
            label="year",
            score=0.93,
            start=106,
            end=110,
        ),
    ]


@pytest.fixture
def mock_ner_entities_low_confidence():
    """NER entities with low confidence scores."""
    from app.schemas import NEREntity

    return [
        NEREntity(
            text="Nguyen", label="name", score=0.35, start=0, end=6
        ),
        NEREntity(
            text="some text",
            label="skills",
            score=0.40,
            start=7,
            end=16,
        ),
    ]


@pytest.fixture
def vietnamese_cv_text():
    """Sample Vietnamese CV text."""
    return """
    NGUYỄN VĂN A
    Email: nguyenvana@gmail.com
    Điện thoại: 0901234567
    Địa chỉ: Quận 1, TP. Hồ Chí Minh

    MỤC TIÊU NGHỀ NGHIỆP
    Mong muốn trở thành lập trình viên Full-stack với kinh nghiệm 3 năm.

    KINH NGHIỆM LÀM VIỆC
    Lập trình viên - Công ty ABC - 2021 đến 2024
    - Phát triển ứng dụng web sử dụng React và Node.js
    - Quản lý cơ sở dữ liệu PostgreSQL

    HỌC VẤN
    Cử nhân Công nghệ Thông tin - Đại học Bách Khoa - 2021

    KỸ NĂNG
    Python, JavaScript, React, Node.js, PostgreSQL, Docker
    """


@pytest.fixture
def english_cv_text():
    """Sample English CV text."""
    return """
    John Doe
    Email: john.doe@email.com
    Phone: +1 555-123-4567
    Location: San Francisco, CA

    EXPERIENCE
    Software Engineer - Google - 2020 to 2024
    - Developed microservices using Java and Spring Boot
    - Led team of 5 engineers

    EDUCATION
    BS Computer Science - MIT - 2020

    SKILLS
    Python, Java, Docker, Kubernetes, AWS, React
    """
