"""Text extraction service for PDF and DOCX files.

Extracts raw text from CV files using:
- pdfplumber for text-based PDFs
- pytesseract (Tesseract OCR) for scanned PDFs
- python-docx for DOCX files (including text boxes)
"""

import io
import logging
import re
import string
from typing import Optional

import pdfplumber
import pytesseract
from PIL import Image
from docx import Document as DocxDocument

from app.config import settings
from app.schemas import TextExtractionResult

logger = logging.getLogger(__name__)


# ── Text quality utilities (inlined from utils/text_quality.py) ────────


def _calculate_printable_ratio(text: str) -> float:
    """Calculate the ratio of printable/valid characters in text."""
    if not text:
        return 0.0
    valid_count = 0
    for char in text:
        if char in string.printable:
            valid_count += 1
        elif char.isalpha():
            valid_count += 1
        elif char.isdigit():
            valid_count += 1
    return valid_count / len(text)


def _is_text_garbled(text: str, threshold: float = 0.70) -> bool:
    """Check if extracted text is likely garbled/corrupted."""
    if not text or len(text) < 10:
        return False
    return _calculate_printable_ratio(text) < threshold


def clean_extracted_text(text: str) -> str:
    """Clean up raw extracted text.

    - Collapse multiple whitespace/newlines
    - Strip leading/trailing whitespace
    - Remove null bytes
    """
    if not text:
        return ""
    text = text.replace("\x00", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[^\S\n]+", " ", text)
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    return text.strip()


# ── Word XML namespaces for text box extraction ────────────────────────

DOCX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "v": "urn:schemas-microsoft-com:vml",
}


# ── Main extraction function ──────────────────────────────────────────


def extract_text(file_content: bytes, mime_type: str) -> TextExtractionResult:
    """Extract text from a CV file.

    Args:
        file_content: Raw bytes of the file.
        mime_type: Detected MIME type (from file_validator).

    Returns:
        TextExtractionResult with extracted text, method used, and warnings.
    """
    if mime_type == "application/pdf":
        return _extract_from_pdf(file_content)
    elif mime_type.endswith("wordprocessingml.document"):
        return _extract_from_docx(file_content)
    else:
        return TextExtractionResult(
            warnings=[f"Unsupported MIME type for text extraction: {mime_type}"]
        )


def _extract_from_pdf(file_content: bytes) -> TextExtractionResult:
    """Extract text from PDF, with per-page OCR fallback."""
    result = TextExtractionResult(method="pdfplumber")
    page_texts: list[str] = []
    ocr_used = False

    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                page_num = page_idx + 1

                # Try text extraction first
                try:
                    text = page.extract_text() or ""
                except Exception as e:
                    logger.warning(
                        "pdfplumber failed on page %d: %s", page_num, e
                    )
                    text = ""

                text = clean_extracted_text(text)

                # Check if text is sufficient
                if len(text) < settings.MIN_TEXT_LENGTH:
                    logger.info(
                        "Page %d has insufficient text (%d chars), "
                        "attempting OCR...",
                        page_num,
                        len(text),
                    )
                    ocr_text = _ocr_pdf_page(page, page_num)
                    if ocr_text and len(ocr_text) > len(text):
                        text = ocr_text
                        ocr_used = True
                        result.warnings.append(
                            f"page_{page_num}_extracted_via_ocr"
                        )

                # Check for garbled text (font encoding issues)
                if text and _is_text_garbled(text):
                    logger.warning(
                        "Page %d has garbled text (low printable ratio), "
                        "attempting OCR...",
                        page_num,
                    )
                    ocr_text = _ocr_pdf_page(page, page_num)
                    if ocr_text and not _is_text_garbled(ocr_text):
                        text = ocr_text
                        ocr_used = True
                        result.warnings.append(
                            f"page_{page_num}_garbled_text_replaced_by_ocr"
                        )
                    else:
                        result.warnings.append(
                            f"page_{page_num}_garbled_text_detected"
                        )

                page_texts.append(text)

    except Exception as e:
        logger.error("PDF text extraction failed: %s", e)
        result.warnings.append(f"pdf_extraction_error: {str(e)}")
        return result

    result.page_texts = page_texts
    result.text = clean_extracted_text("\n\n".join(page_texts))
    result.ocr_used = ocr_used

    if ocr_used:
        result.method = "pdfplumber+ocr_tesseract"

    # Final check: if total text is still too short
    if len(result.text) < settings.MIN_TEXT_LENGTH:
        result.warnings.append("possible_scanned_file")

    logger.info(
        "PDF extraction complete: %d pages, %d chars, method=%s, ocr=%s",
        len(page_texts),
        len(result.text),
        result.method,
        ocr_used,
    )

    return result


def _ocr_pdf_page(page, page_num: int) -> Optional[str]:
    """Perform OCR on a single PDF page using Tesseract."""
    try:
        img = page.to_image(resolution=300)
        pil_image = img.original

        ocr_text = pytesseract.image_to_string(
            pil_image, lang=settings.TESSERACT_LANG
        )
        ocr_text = clean_extracted_text(ocr_text)

        logger.info(
            "OCR on page %d: extracted %d chars", page_num, len(ocr_text)
        )
        return ocr_text

    except Exception as e:
        logger.warning("OCR failed on page %d: %s", page_num, e)
        return None


def _extract_from_docx(file_content: bytes) -> TextExtractionResult:
    """Extract text from DOCX, including text boxes."""
    result = TextExtractionResult(method="python_docx")
    parts: list[str] = []

    try:
        doc = DocxDocument(io.BytesIO(file_content))

        # 1. Extract from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 2. Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        # 3. Extract from text boxes (often missed by python-docx)
        textbox_texts = _extract_docx_textboxes(doc)
        if textbox_texts:
            parts.extend(textbox_texts)
            result.warnings.append(
                f"extracted_{len(textbox_texts)}_text_boxes"
            )

    except Exception as e:
        logger.error("DOCX text extraction failed: %s", e)
        result.warnings.append(f"docx_extraction_error: {str(e)}")
        return result

    result.text = clean_extracted_text("\n".join(parts))
    result.page_texts = [result.text]  # DOCX doesn't have clear page breaks

    if len(result.text) < settings.MIN_TEXT_LENGTH:
        result.warnings.append("very_short_docx_content")

    logger.info(
        "DOCX extraction complete: %d chars, %d text boxes found",
        len(result.text),
        len(textbox_texts) if textbox_texts else 0,
    )

    return result


def _extract_docx_textboxes(doc: DocxDocument) -> list[str]:
    """Extract text from DOCX text boxes and shapes."""
    texts: list[str] = []

    try:
        body = doc.element.body

        # Find text in WordprocessingML shapes (wps:txbx)
        for txbx in body.iter(
            "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"
        ):
            for p in txbx.iter(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
            ):
                if p.text and p.text.strip():
                    texts.append(p.text.strip())

        # Find text in VML shapes (v:textbox)
        for textbox in body.iter(
            "{urn:schemas-microsoft-com:vml}textbox"
        ):
            for p in textbox.iter(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
            ):
                if p.text and p.text.strip():
                    texts.append(p.text.strip())

    except Exception as e:
        logger.warning("Text box extraction failed: %s", e)

    return texts
